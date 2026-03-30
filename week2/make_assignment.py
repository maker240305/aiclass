from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_korean_font(run, size=11):
    run.font.size = Pt(size)
    run.font.name = 'Malgun Gothic'
    # Important: Set the actual font using rFonts for Korean characters
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def create_assignment():
    doc = Document()
    
    # Title
    title_text = '[AI 활용 수업 실습 및 과제 1]'
    title = doc.add_heading('', 0)
    run = title.add_run(title_text)
    set_korean_font(run, 24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Student Info
    p = doc.add_paragraph()
    
    run = p.add_run('성명: ')
    run.bold = True
    set_korean_font(run)
    run = p.add_run('김철수 (신입생)\n')
    set_korean_font(run)
    
    run = p.add_run('학번: ')
    run.bold = True
    set_korean_font(run)
    run = p.add_run('20260316\n')
    set_korean_font(run)
    
    run = p.add_run('전공: ')
    run.bold = True
    set_korean_font(run)
    run = p.add_run('컴퓨터공학과\n')
    set_korean_font(run)
    
    run = p.add_run('장래희망: ')
    run.bold = True
    set_korean_font(run)
    run = p.add_run('혁신적인 게임 개발자')
    set_korean_font(run)
    
    doc.add_paragraph('\n')
    
    # Section 1
    doc.add_heading('1. 컴퓨터에 대한 나의 생각과 막연함', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '신입생으로서 컴퓨터공학은 즐거워 보이지만 한편으로는 두렵기도 한 분야였습니다. '
        '특히 "개발 환경 구축"이라는 첫 단추부터 파이썬 버전 충돌이나 환경 변수 설정 같은 '
        '복잡한 문제들을 마주했을 때, 누구의 도움 없이는 시작조차 못 할 것 같아 막막함을 느꼈습니다.'
    )
    set_korean_font(run)
    
    # Section 2
    doc.add_heading('2. AI Agent를 통해 해결한 일과 변화된 생각', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '하지만 오늘 수업을 통해 AI Agent를 활용해 보면서 생각이 완전히 바뀌었습니다! '
        '파이썬 3.14.3 최신 버전을 설치하고, 윈도우 시스템 변수 순서 때문에 발생한 꼬인 문제들을 '
        'AI Agent와 대화하며 하나씩 풀어나갈 수 있었습니다. 예전에는 몇 시간을 검색해도 몰랐을 내용들을 '
        '옆에서 과외 선생님이 알려주듯 설명해주니, 이제는 어려운 기술적 문제도 겁내지 않고 도전할 수 있겠다는 자신감이 생겼습니다.'
    )
    set_korean_font(run)
    
    # Section 3
    doc.add_heading('3. AI Agent와 함께 당장 하고 싶은 일들', level=1)
    
    tasks = [
        ('나만의 인공지능 NPC 설계', '플레이어와 자유롭게 대화할 수 있는 NPC의 대사 로직을 AI Agent와 함께 설계하고 싶습니다.'),
        ('게임 프로토타입 코딩', '복잡한 물리 공식이나 렌더링 코드를 직접 다 짜기 전에, AI Agent의 도움을 받아 핵심 게임성만 담은 최소 기능 제품(MVP)을 빠르게 만들어보고 싶습니다.'),
        ('1인 개발 프로젝트 가동', '스토리 기획, 배경 음악, 캐릭터 디자인 등 제가 부족한 부분들을 AI Agent를 활용해 보완함으로써, 혼자서도 하나의 완성된 게임 세상을 창조해보고 싶습니다.')
    ]
    
    for t_title, t_desc in tasks:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{t_title}: ')
        run.bold = True
        set_korean_font(run)
        run = p.add_run(t_desc)
        set_korean_font(run)
        
    doc.add_paragraph('\n')
    
    # Conclusion
    doc.add_heading('마치며', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        'AI Agent는 이제 저에게 단순한 소프트웨어가 아니라, 함께 성장해 나갈 든든한 "사수"와 같습니다. '
        '이 수업을 통해 인공지능을 완벽히 활용하는 법을 익혀, 전 세계 사람들을 놀라게 할 멋진 게임을 만드는 개발자가 되겠습니다!'
    )
    set_korean_font(run)
    
    # Save document
    current_dir = os.path.dirname(__file__)
    file_name = 'AI_활용_과제_게임개발자_김철수.docx'
    file_path = os.path.join(current_dir, file_name)
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    result_path = create_assignment()
    print(f"File created at: {result_path}")
