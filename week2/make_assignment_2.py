from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_korean_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.name = 'Malgun Gothic'
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def create_assignment_2():
    doc = Document()
    
    # Title
    title_text = '[AI 활용 수업 실습 및 과제 2]'
    title = doc.add_heading('', 0)
    run = title.add_run(title_text)
    set_korean_font(run, 24, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Concept
    doc.add_heading('1. 대학생활 컨셉: "몰입하는 크리에이티브 아키텍트"', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '저의 대학생활 컨셉은 단순히 코드를 짜는 사람이 아닌, "세상의 재미를 설계하는 아키텍트"입니다. '
        'AI 기술을 도구 삼아 상상을 현실로 구현하는 창의적인 개발 환경에 스스로를 몰입시키는 것이 목표입니다.'
    )
    set_korean_font(run)

    # Section 2: Goals & Strategy
    doc.add_heading('2. 성공의 목적과 실행 전략', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    run = hdr_cells[0].paragraphs[0].add_run('성공의 목적 (Goal)')
    set_korean_font(run, bold=True)
    run = hdr_cells[1].paragraphs[0].add_run('실행 방법 (Execution)')
    set_korean_font(run, bold=True)
    
    data = [
        ('기술적 자립 (Technical Mastery)', '매일 1시간 AI Agent와 함께 고난도 알고리즘 및 게임 엔진(Unity/Unreal) 학습'),
        ('포트폴리오 구축 (Build)', '매 학기 1개의 소규모 인디 게임 프로젝트 완성 및 GitHub 기록'),
        ('네트워킹 (Connection)', '개발자 커뮤니티 활동 및 AI 컨퍼런스 참여로 현업 전문가들과 소통')
    ]
    
    for goal, exe in data:
        row_cells = table.add_row().cells
        run_g = row_cells[0].paragraphs[0].add_run(goal)
        set_korean_font(run_g)
        run_e = row_cells[1].paragraphs[0].add_run(exe)
        set_korean_font(run_e)

    doc.add_paragraph('\n')

    # Section 3: Visualization (Diagram)
    doc.add_heading('3. 대학생활 성공 로드맵 도식화', level=1)
    p = doc.add_paragraph()
    run = p.add_run('AI를 활용하여 생성한 저의 성장 로드맵 컨셉 아트입니다.')
    set_korean_font(run)
    
    # Image path (Using the newly generated AI image for high visual quality)
    image_path = r'C:\Users\john0\.gemini\antigravity\brain\a0a6a81e-b9a7-4f48-acac-2ef2eba0354a\nanobanana_concept_diagram_1773641520877.png'
    try:
        doc.add_picture(image_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f'(이미지 로드 실패: {e})')

    # Footer
    doc.add_paragraph('\n')
    p = doc.add_paragraph()
    run = p.add_run('※ 위 도식은 AI 프롬프트 엔지니어링을 통해 저의 비전을 시각화한 결과물입니다.')
    set_korean_font(run, size=9)

    # Save document
    current_dir = os.path.dirname(__file__)
    file_name = 'AI_활용_과제_2_성공로드맵_김철수.docx'
    file_path = os.path.join(current_dir, file_name)
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    result_path = create_assignment_2()
    print(f"File created at: {result_path}")
