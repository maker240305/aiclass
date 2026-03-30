from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_korean_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.name = 'Malgun Gothic'
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def create_integrated_assignment():
    doc = Document()
    
    # -----------------------------
    # 메인 타이틀 및 학생 정보
    # -----------------------------
    title_text = '[AI 활용 수업 종합 실습 및 과제]'
    title = doc.add_heading('', 0)
    run = title.add_run(title_text)
    set_korean_font(run, 26, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 학생 정보
    info_lines = [
        ('소속', '컴퓨터공학과 (26학번)'),
        ('성명', '김철수 (신입생)'),
        ('학번', '20260316'),
        ('장래희망', 'AI 풀스택 웹/게임 개발자')
    ]
    for key, val in info_lines:
        run = p.add_run(f'{key}: ')
        set_korean_font(run, bold=True)
        run = p.add_run(f'{val}\n')
        set_korean_font(run)
        
    doc.add_page_break()

    # -----------------------------
    # [과제 1] 컴퓨터에 대한 나의 생각, 막연함, 그리고 AI Agent의 도움
    # -----------------------------
    doc.add_heading('[과제 1] 컴퓨터에 대한 막연함과 AI Agent의 활용', level=1)
    
    doc.add_heading('1. 컴퓨터에 대한 초기 막연함', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '신입생으로서 컴퓨터공학은 매력적이었지만 한편으로는 두렵기도 한 분야였습니다. '
        '초기 파이썬 개발 환경 구축 시 버전 충돌이나 시스템/사용자 환경 변수 세팅 같은 복잡한 문제들을 마주했을 때, '
        '도움 없이는 시작하기 힘들다는 막막함을 느꼈습니다.'
    )
    set_korean_font(run)

    doc.add_heading('2. AI Agent를 통해 해결한 일련의 과정', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '이번 수업에서 AI Agent와 대화하며, 파이썬 최신 버전을 설치하고 윈도우 환경 변수의 복잡성을 해결했습니다. '
        '제가 이해하기 쉽게 에러의 원인과 윈도우 OS의 특성(시스템/유저 변수)을 알려주는 AI의 과정은 단순한 검색 그 이상의 학습 경험이었습니다.'
    )
    set_korean_font(run)

    doc.add_heading('3. AI Agent와 함께 당장 하고 싶은 일들', level=2)
    tasks1 = [
        ('지능형 NPC 대사 시스템', '플레이어와 맥락에 맞춰 자유롭게 대화할 수 있는 AI NPC 로직 설계.'),
        ('게임 프로토타입 구현', '물리엔진 등 복잡한 코어 로직을 직접 짜기 전, AI의 도움을 받아 게임의 핵심 메커니즘만 빠르게 구현(MVP).'),
        ('1인 개발 프로젝트 가동', '제가 부족한 배경/캐릭터/스토리라인 기획을 AI를 통해 보완하며 1인 개발 완성하기.')
    ]
    for t_title, t_desc in tasks1:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{t_title}: ')
        set_korean_font(run, bold=True)
        run = p.add_run(t_desc)
        set_korean_font(run)

    doc.add_page_break()

    # -----------------------------
    # [과제 2] 대학생활 성공 로드맵 도식화
    # -----------------------------
    doc.add_heading('[과제 2] 대학생활 컨셉 및 성공 로드맵', level=1)
    
    doc.add_heading('1. 대학생활 컨셉: "몰입하는 크리에이티브 아키텍트"', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '단순한 코더를 넘어, AI라는 강력한 도구로 게임/서비스 세상의 질서와 재미를 설계하고 실현하는 아키텍트가 되는 것입니다.'
    )
    set_korean_font(run)

    doc.add_heading('2. 대학생활 성공 로드맵 도식화 (AI 이미지 생성 활용)', level=2)
    p = doc.add_paragraph()
    run = p.add_run('가운데 핵심 컨셉을 두고 3개의 목적이 집중되는 AI 도식화 아트워크입니다.')
    set_korean_font(run)
    
    image_path = r'C:\Users\john0\.gemini\antigravity\brain\a0a6a81e-b9a7-4f48-acac-2ef2eba0354a\nanobanana_concept_diagram_1773641520877.png'
    try:
        doc.add_picture(image_path, width=Inches(6.0))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f'(이미지 로드 실패: {e})')

    doc.add_heading('3. 성공 목적 및 일상 실행 전략', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    run = hdr_cells[0].paragraphs[0].add_run('목적 (Goal)')
    set_korean_font(run, bold=True)
    run = hdr_cells[1].paragraphs[0].add_run('실행 방법 (Execution)')
    set_korean_font(run, bold=True)
    
    data2 = [
        ('기술적 자립 (Technical Mastery)', '매일 1시간 AI Agent와 함께 고난도 알고리즘/엔진 학습'),
        ('포트폴리오 구축 (Build)', '매 학기 1개 규모의 소규모 웹/게임 프로젝트 런칭 및 GitHub 기록'),
        ('글로벌 네트워킹 (Connection)', '국내외 오픈소스 커뮤니티 참여 및 컨퍼런스 기여')
    ]
    for goal, exe in data2:
        row_cells = table.add_row().cells
        run_g = row_cells[0].paragraphs[0].add_run(goal)
        set_korean_font(run_g)
        run_e = row_cells[1].paragraphs[0].add_run(exe)
        set_korean_font(run_e)

    doc.add_page_break()

    # -----------------------------
    # [과제 3] 효율적인 프롬프팅: 배경지식과 세부 학습 계획
    # -----------------------------
    doc.add_heading('[과제 3] 역량 기반 구체적 4년 학습 계획', level=1)
    
    doc.add_heading('1. 현 상황 및 배경지식 소개', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '저는 파이썬 기본 문법과 로직 구성력, HTML/CSS/JavaScript와 같은 기본 웹 지식을 갖춘 컴퓨터공학과 전공생입니다. '
        '이제는 이 파편화된 기초 역량을 결합하여 체계적인 서비스형 소프트웨어로 탈바꿈시키는 실전 아키텍처 역량이 필요합니다.'
    )
    set_korean_font(run)

    doc.add_heading('2. 대학생활 최종 목표', level=2)
    p = doc.add_paragraph()
    run_bold = p.add_run('"AI 연동 풀스택 프레임워크 구축"')
    set_korean_font(run_bold, bold=True)
    run = p.add_run(
        '\n단순한 파이썬 스크립트 작성에서 벗어나, 제가 만든 AI 로직을 실제 유저들이 상호작용할 수 있는 웹 기반(클라우드 배포) 포트폴리오로 구축하는 기획자가 되겠습니다.'
    )
    set_korean_font(run)

    doc.add_heading('3. 대학생활 4개년 세부 학습 및 실행 계획', level=2)
    plans3 = [
        ('[1학년] 기초 확장 및 결합', '기보유한 파이썬/웹 역량을 통합하기 위해 Django나 FastAPI 등 백엔드 프레임워크 학습 및 미니 게시판 구축.'),
        ('[2학년] AI API 융합', 'OpenAI, Gemini API를 학습하여 1학년때 만든 뼈대 위에 기능으로 활용. 토이 프로젝트 2개 완성.'),
        ('[3학년] 클라우드 실전 배포', 'AWS/Vercel 서버리스 배포를 습득하고 트래픽 관리가 가능한 웹서비스 구축 및 네트워킹 최적화 경험.'),
        ('[4학년] 실무 협업 포트폴리오 완성', '오픈소스/산학연계 프로젝트에서 실무진과 협업하며, 기획부터 런칭까지 혼자 주도할 수 있는 풀스택 최종 작품 포트폴리오 런칭.')
    ]
    for phase, detail in plans3:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{phase}: ')
        set_korean_font(run, bold=True)
        run = p.add_run(detail)
        set_korean_font(run)

    # -----------------------------
    # 종합 결론
    # -----------------------------
    doc.add_paragraph('\n')
    doc.add_heading('[종합 결론]', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '신입생의 두려움은 AI Agent를 "사수" 삼음으로써 크나큰 호기심과 계획성으로 탈바꿈했습니다. '
        '기본 파이썬 지식과 웹 기반 기술을 엮어, 대학생활 4년 동안 저는 세상에서 가장 재미있고 직관적인 기술을 설계하는 크리에이티브 아키텍트로 성장할 것입니다.'
    )
    set_korean_font(run)

    current_dir = os.path.dirname(__file__)
    file_name = 'AI_활용_과제_통합본_김철수.docx'
    file_path = os.path.join(current_dir, file_name)
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    result_path = create_integrated_assignment()
    print(f"File created at: {result_path}")
