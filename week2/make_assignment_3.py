from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def set_korean_font(run, size=11, bold=False):
    run.font.size = Pt(size)
    run.font.name = 'Malgun Gothic'
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def create_assignment_3():
    doc = Document()
    
    # Title
    title_text = '[AI 활용 수업 실습 및 과제 3]'
    title = doc.add_heading('', 0)
    run = title.add_run(title_text)
    set_korean_font(run, 24, True)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Current Status & Background
    doc.add_heading('1. 현 상황 및 배경지식 소개', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '저는 컴퓨터공학과에 재학 중인 전공자로서, 현재 두 가지 핵심 기초 역량을 갖추고 있습니다.\n'
        '첫째, 파이썬(Python) 기본 문법과 데이터 처리 로직을 이해하고 활용할 수 있습니다.\n'
        '둘째, HTML/CSS/JavaScript를 아우르는 기초적인 웹 프로그래밍 환경을 경험해 보았습니다.\n'
        '이제 이 파편화된 기초 지식들을 결합하여, AI 기술을 실제 웹 서비스나 게임 백엔드로 구현하는 실전 능력이 필요한 단계입니다.'
    )
    set_korean_font(run)

    # Section 2: Specific Goals
    doc.add_heading('2. 구체적인 목표 설정', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '대학생활 동안 가장 달성하고 싶은 최종 목표는 '
    )
    set_korean_font(run)
    run_bold = p.add_run('"AI 연동 풀스택 웹/게임 서비스 개발 프레임워크 구축"')
    set_korean_font(run_bold, bold=True)
    run = p.add_run(
        '입니다. 단순히 코드를 짜는 것을 넘어, 제가 만든 파이썬 기반 AI 로직을 웹 프로그래밍 지식을 통해 웹상에 시각적으로 배포하고, '
        '실제 유저들이 상호작용할 수 있는 포트폴리오를 만들고 싶습니다.'
    )
    set_korean_font(run)

    # Section 3: Learning Plan & Execution
    doc.add_heading('3. 대학생활 학습계획 및 실행 방법', level=1)
    
    plans = [
        ('1학년 (기초 확장)', '현재 알고 있는 파이썬(Python)과 웹 지식을 심화하기 위해, Django/FastAPI 개발 프레임워크를 학습하여 웹과 파이썬 백엔드를 연결하는 실습을 합니다.'),
        ('2학년 (AI 융합)', '학교에서 배우는 자료구조/알고리즘 외에 개인적으로 AI Agent API(OpenAI, Gemini 등)를 백엔드에 접목하는 토이 프로젝트 2개 완성하기.'),
        ('3학년 (실전 배포)', 'AWS, Vercel 등의 클라우드 환경에 웹 서비스를 직접 배포하고, 100명 이상의 유저 트래픽을 처리해 보는 네트워킹 및 서버 최적화 경험 쌓기.'),
        ('4학년 (포트폴리오 완성)', '오픈소스 커뮤니티 기여 및 산학협력 프로젝트를 통해 실무진과의 협업 능력을 배양하고 최종 풀스택 게임/서비스 포트폴리오를 런칭하기.')
    ]
    
    for grade, detail in plans:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'[{grade}] ')
        set_korean_font(run, bold=True)
        run = p.add_run(detail)
        set_korean_font(run)

    doc.add_paragraph('\n')

    # Conclusion
    doc.add_heading('마치며', level=1)
    p = doc.add_paragraph()
    run = p.add_run(
        '파이썬과 기초 웹 프로그래밍이라는 저의 무기에 AI라는 강력한 엔진을 달아, 4년 뒤에는 설계와 구현을 모두 주도하는 개발자로 성장하겠습니다.'
    )
    set_korean_font(run)

    # Save document
    current_dir = os.path.dirname(__file__)
    file_name = 'AI_활용_과제_3_학습계획_김철수.docx'
    file_path = os.path.join(current_dir, file_name)
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    result_path = create_assignment_3()
    print(f"File created at: {result_path}")
